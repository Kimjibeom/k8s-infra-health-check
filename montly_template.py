import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx 라이브러리가 필요합니다. 'pip install python-docx'를 실행해주세요.")
    exit()

def create_monthly_report_template(output_filename="CMP_Monthly_Report_Form.docx"):
    doc = Document()
    
    # 스타일 설정 (한글 폰트 적용)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # 0. 제목 및 결재란
    title = doc.add_heading(f'{datetime.now().year}년 {datetime.now().month}월 CMP 인프라 운영 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 보고서 개요
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    info_table.rows[0].cells[0].text = "작성일"
    info_table.rows[0].cells[1].text = datetime.now().strftime('%Y년 %m월 %d일')
    info_table.rows[1].cells[0].text = "작성자"
    info_table.rows[1].cells[1].text = "플랫폼팀 (현장대리인: OOO)"
    info_table.rows[2].cells[0].text = "수신"
    info_table.rows[2].cells[1].text = "OOO 귀하"
    
    doc.add_paragraph() 

    # 공통 함수: 테이블 헤더 스타일 적용
    def set_header_style(row):
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            
            # 셀 배경색 설정 (XML 조작 필요, 여기선 생략하고 텍스트만 처리)
            # 폰트만 볼드체로

    # 공통 함수: 빈 행 추가 (기본 4행)
    def add_empty_rows(table, count=4):
        cols = len(table.columns)
        for _ in range(count):
            table.add_row()

    # 1. 주요 운영 현황
    doc.add_heading('1. 주요 운영 현황', level=1)
    doc.add_paragraph("■ 총평: 금월 CMP 인프라 및 서비스는 전반적으로 안정적으로 운영되었습니다. (특이사항 없음)")
    doc.add_paragraph("■ 가동률: 99.9% (주요 서비스 기준)")
    doc.add_paragraph("■ 정기점검 요약:")
    doc.add_paragraph("   - 총 점검 항목: 95개")
    doc.add_paragraph("   - ✅ 정상: 00 / ⚠️ 경고: 00 / ❌ 위험: 00")

    doc.add_paragraph()

    # 2. 주요 작업 및 특이사항
    doc.add_heading('2. 주요 작업 및 특이사항', level=1)
    table2 = doc.add_table(rows=2, cols=3)
    table2.style = 'Table Grid'
    
    # 헤더
    hdr2 = table2.rows[0]
    hdr2.cells[0].text = "구분 (클러스터/담당자)"
    hdr2.cells[1].text = "내용"
    hdr2.cells[2].text = "비고"
    set_header_style(hdr2)
    
    # 예시 데이터 (1행)
    table2.rows[1].cells[0].text = "DEV / 홍길동"
    table2.rows[1].cells[1].text = "ArgoCD 버전 패치 (v2.8 -> v2.9)"
    table2.rows[1].cells[2].text = "완료"
    
    # 빈 행 추가 (4행)
    add_empty_rows(table2, 4)

    doc.add_paragraph()

    # 3. 장애처리 현황
    doc.add_heading('3. 장애처리 현황', level=1)
    
    doc.add_heading('3-1. 장애 발생/처리 현황', level=2)
    table3_1 = doc.add_table(rows=2, cols=6)
    table3_1.style = 'Table Grid'
    
    hdr3_1 = table3_1.rows[0]
    headers3_1 = ["대상시스템", "장애일시", "장애현황", "장애원인", "조치결과", "비고"]
    for i, h in enumerate(headers3_1):
        hdr3_1.cells[i].text = h
    set_header_style(hdr3_1)
    
    # 예시 데이터
    row = table3_1.rows[1]
    row.cells[0].text = "RabbitMQ"
    row.cells[1].text = "2026-01-05 14:00"
    row.cells[2].text = "Pod CrashLoop"
    row.cells[3].text = "OOM (메모리 부족)"
    row.cells[4].text = "Limit 증설 후 재기동"
    row.cells[5].text = "정상화"
    
    add_empty_rows(table3_1, 4)

    doc.add_paragraph()

    doc.add_heading('3-2. 장애 보고서 첨부 (통계)', level=2)
    table3_2 = doc.add_table(rows=2, cols=5)
    table3_2.style = 'Table Grid'
    
    hdr3_2 = table3_2.rows[0]
    headers3_2 = ["유형", "전월 장애 건수", "금월 장애 건수", "서비스 중단 시간", "증감"]
    for i, h in enumerate(headers3_2):
        hdr3_2.cells[i].text = h
    set_header_style(hdr3_2)
    
    # 예시 데이터
    row = table3_2.rows[1]
    row.cells[0].text = "Critical"
    row.cells[1].text = "1"
    row.cells[2].text = "0"
    row.cells[3].text = "0분"
    row.cells[4].text = "▼ 1"
    
    add_empty_rows(table3_2, 4)

    doc.add_paragraph()

    # 4. 시스템별 운영 현황
    doc.add_heading('4. 시스템별 운영 현황', level=1)

    doc.add_heading('4-1. Container 운영 내역', level=2)
    table4_1 = doc.add_table(rows=2, cols=6)
    table4_1.style = 'Table Grid'
    
    hdr4_1 = table4_1.rows[0]
    headers4_1 = ["클러스터", "네임스페이스", "워크로드명", "Replica", "이미지 버전", "상태"]
    for i, h in enumerate(headers4_1):
        hdr4_1.cells[i].text = h
    set_header_style(hdr4_1)
    
    # 예시 데이터
    row = table4_1.rows[1]
    row.cells[0].text = "PRD"
    row.cells[1].text = "argocd"
    row.cells[2].text = "argocd-server"
    row.cells[3].text = "2 / 2"
    row.cells[4].text = "v2.9.3"
    row.cells[5].text = "정상"
    
    add_empty_rows(table4_1, 4)

    doc.add_paragraph()

    doc.add_heading('4-2. 스토리지 운영 내역 (변경 관리)', level=2)
    table4_2 = doc.add_table(rows=2, cols=5)
    table4_2.style = 'Table Grid'
    
    hdr4_2 = table4_2.rows[0]
    headers4_2 = ["일자", "요청자", "내용", "대상 PVC/Volume", "비고"]
    for i, h in enumerate(headers4_2):
        hdr4_2.cells[i].text = h
    set_header_style(hdr4_2)
    
    # 예시 데이터
    row = table4_2.rows[1]
    row.cells[0].text = "2026-01-02"
    row.cells[1].text = "김개발"
    row.cells[2].text = "로그 볼륨 증설 (10G->20G)"
    row.cells[3].text = "loki-store-pvc"
    row.cells[4].text = "완료"
    
    add_empty_rows(table4_2, 4)

    doc.add_paragraph()

    doc.add_heading('4-3. 스토리지 사용 현황 (Capacity)', level=2)
    table4_3 = doc.add_table(rows=2, cols=6)
    table4_3.style = 'Table Grid'
    
    hdr4_3 = table4_3.rows[0]
    headers4_3 = ["구분", "네임스페이스", "PVC 명", "할당 용량", "상태", "연결 파드"]
    for i, h in enumerate(headers4_3):
        hdr4_3.cells[i].text = h
    set_header_style(hdr4_3)
    
    # 예시 데이터
    row = table4_3.rows[1]
    row.cells[0].text = "Block"
    row.cells[1].text = "databases"
    row.cells[2].text = "data-mariadb-0"
    row.cells[3].text = "8Gi"
    row.cells[4].text = "Bound"
    row.cells[5].text = "mariadb-0"
    
    add_empty_rows(table4_3, 4)

    doc.add_paragraph()

    # 5. SSL 인증서 관리
    doc.add_heading('5. SSL 인증서 관리', level=1)
    table5 = doc.add_table(rows=2, cols=6)
    table5.style = 'Table Grid'
    
    hdr5 = table5.rows[0]
    headers5 = ["도메인", "인증서 발급기관", "만료일자", "잔여일수", "상태", "비고"]
    for i, h in enumerate(headers5):
        hdr5.cells[i].text = h
    set_header_style(hdr5)
    
    # 예시 데이터
    row = table5.rows[1]
    row.cells[0].text = "google.com"
    row.cells[1].text = "GTS CA 1C3"
    row.cells[2].text = "2026-02-25"
    row.cells[3].text = "51일"
    row.cells[4].text = "경고"
    row.cells[5].text = "갱신 예정"
    
    add_empty_rows(table5, 4)

    doc.add_paragraph()

    # 6. 오픈소스 SW 운영 현황
    doc.add_heading('6. 오픈소스 SW 운영 현황', level=1)
    table6 = doc.add_table(rows=2, cols=6)
    table6.style = 'Table Grid'
    
    hdr6 = table6.rows[0]
    headers6 = ["분류", "SW 명", "현재 버전", "최신 버전", "최근 업데이트", "비고"]
    for i, h in enumerate(headers6):
        hdr6.cells[i].text = h
    set_header_style(hdr6)
    
    # 예시 데이터
    row = table6.rows[1]
    row.cells[0].text = "Runtime"
    row.cells[1].text = "Containerd"
    row.cells[2].text = "v1.6.28"
    row.cells[3].text = "v1.7.13"
    row.cells[4].text = "2025-12-20"
    row.cells[5].text = "-"
    
    add_empty_rows(table6, 4)

    doc.add_paragraph()

    # 7. 알림 현황 관리
    doc.add_heading('7. 알림(Alert) 현황 관리', level=1)
    table7 = doc.add_table(rows=2, cols=6)
    table7.style = 'Table Grid'
    
    hdr7 = table7.rows[0]
    headers7 = ["분류", "임계치 기준", "기준 값", "발생 수", "증감", "대상"]
    for i, h in enumerate(headers7):
        hdr7.cells[i].text = h
    set_header_style(hdr7)
    
    # 예시 데이터
    row = table7.rows[1]
    row.cells[0].text = "Node"
    row.cells[1].text = "CPU 사용률"
    row.cells[2].text = "> 90%"
    row.cells[3].text = "2건"
    row.cells[4].text = "▼ 1"
    row.cells[5].text = "Worker #1"
    
    add_empty_rows(table7, 4)

    doc.add_paragraph()

    # 8. 인원 상세 현황
    doc.add_heading('8. 인원 상세 현황', level=1)
    table8 = doc.add_table(rows=2, cols=5)
    table8.style = 'Table Grid'
    
    hdr8 = table8.rows[0]
    headers8 = ["소속", "성명", "직급/역할", "연락처", "비상연락"]
    for i, h in enumerate(headers8):
        hdr8.cells[i].text = h
    set_header_style(hdr8)
    
    # 예시 데이터
    row = table8.rows[1]
    row.cells[0].text = "플랫폼팀"
    row.cells[1].text = "홍길동"
    row.cells[2].text = "책임/PL"
    row.cells[3].text = "010-1234-5678"
    row.cells[4].text = "O"
    
    add_empty_rows(table8, 4)

    # 파일 저장
    doc.save(output_filename)
    print(f"✅ 보고서 양식 생성 완료: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    create_monthly_report_template()